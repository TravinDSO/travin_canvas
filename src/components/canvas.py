"""
Markdown Canvas Component for Travin Canvas

This module provides the canvas component for the Travin Canvas application.
It implements a collaborative markdown editor with real-time preview,
formatting tools, and document management features to support
document drafting and refinement.

Key features:
- Split-view markdown editor with real-time preview
- Document enhancement tools powered by LLM
- Document history with undo capability
- File import and export functionality
- Table of contents generation

Dependencies:
- streamlit: For UI components
- utils.markdown_utils: For markdown processing
- utils.llm_utils: For LLM-powered enhancements
"""

import os
import json
import streamlit as st
from utils.markdown_utils import MarkdownProcessor
from utils.llm_utils import LLMManager
import docx
from PyPDF2 import PdfReader
import io

class MarkdownCanvas:
    """
    Implements the markdown canvas for the Travin Canvas application.
    
    This class manages the document editing experience, including:
    - Markdown editing with syntax highlighting
    - Real-time preview rendering with toggle view
    - Document history management with undo capability
    - Streamlined file operations with unified toolbar
    
    The canvas is rendered in the main content area and maintains
    its state between application reruns using session state.
    """
    
    def __init__(self, on_content_change=None):
        """
        Initialize the markdown canvas.
        
        Args:
            on_content_change (callable, optional): Callback for content changes
        """
        self.markdown_processor = MarkdownProcessor()
        self.llm_manager = LLMManager()
        self.on_content_change = on_content_change
        
        # Initialize session state variables
        if "markdown_content" not in st.session_state:
            st.session_state.markdown_content = ""
            
        if "view_mode" not in st.session_state:
            st.session_state.view_mode = "editor"
            
        if "current_file" not in st.session_state:
            st.session_state.current_file = None
            
        # Initialize undo history
        if "undo_history" not in st.session_state:
            st.session_state.undo_history = []
            
        if "max_history" not in st.session_state:
            st.session_state.max_history = 10
            
        # Add a flag to track file processing status
        if "file_processed" not in st.session_state:
            st.session_state.file_processed = False
            
        if "last_uploaded_file" not in st.session_state:
            st.session_state.last_uploaded_file = None
    
    def render(self):
        """Render the markdown canvas."""
        # Minimal CSS for basic layout
        st.markdown("""
        <style>
        .stButton button {
            font-weight: 500;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Display current file name if available
        if st.session_state.current_file:
            st.caption(f"Current file: {st.session_state.current_file}")
        
        # Combined toolbar with all controls in a single row
        col1, col2, col3, col4, col5 = st.columns([1, 1, 2, 1, 1])
        
        # New document button
        with col1:
            if st.button("📄 New", use_container_width=True):
                # Save current state to history before clearing
                self._save_to_history()
                st.session_state.markdown_content = ""
                st.session_state.current_file = None
                st.rerun()
        
        # Undo button
        with col2:
            if st.button("↩️ Undo", disabled=len(st.session_state.undo_history) == 0, use_container_width=True):
                self._undo_last_change()
        
        # File uploader
        with col3:
            uploaded_file = st.file_uploader("Upload Document", type=["md", "txt", "docx", "pdf"], label_visibility="collapsed")
            
            # Process the file only if it's a new file or hasn't been processed yet
            if uploaded_file and (uploaded_file != st.session_state.last_uploaded_file or not st.session_state.file_processed):
                try:
                    # Save current state to history before loading new file
                    self._save_to_history()
                    
                    # Update the last uploaded file
                    st.session_state.last_uploaded_file = uploaded_file
                    st.session_state.file_processed = False
                    
                    # Process the file based on its type
                    file_extension = uploaded_file.name.split('.')[-1].lower()
                    
                    # Add a status message
                    with st.status(f"Processing {file_extension} file: {uploaded_file.name}...", expanded=True) as status:
                        if file_extension in ['md', 'txt']:
                            # Handle markdown and text files
                            content = uploaded_file.getvalue().decode("utf-8")
                            st.session_state.markdown_content = content
                            status.update(label=f"Loaded text file: {len(content)} characters", state="complete")
                        elif file_extension == 'docx':
                            # Handle Word documents
                            content = self._extract_text_from_docx(uploaded_file)
                            st.session_state.markdown_content = content
                            status.update(label=f"Converted Word document: {len(content)} characters", state="complete")
                        elif file_extension == 'pdf':
                            # Handle PDF files
                            content = self._extract_text_from_pdf(uploaded_file)
                            st.session_state.markdown_content = content
                            status.update(label=f"Converted PDF file: {len(content)} characters", state="complete")
                        
                        st.session_state.current_file = uploaded_file.name
                        
                        # Notify about content change
                        if self.on_content_change:
                            self.on_content_change(st.session_state.markdown_content)
                        
                        # Mark as processed to avoid reprocessing
                        st.session_state.file_processed = True
                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")
                    import traceback
                    st.error(traceback.format_exc())
        
        # Toggle between editor and preview modes
        with col4:
            editor_button_style = "primary" if st.session_state.view_mode == "editor" else "secondary"
            if st.button("✏️ Edit", type=editor_button_style, use_container_width=True):
                st.session_state.view_mode = "editor"
                st.rerun()
        
        with col5:
            preview_button_style = "primary" if st.session_state.view_mode == "preview" else "secondary"
            if st.button("👁️ View", type=preview_button_style, use_container_width=True):
                st.session_state.view_mode = "preview"
                st.rerun()
        
        # Render based on selected view mode
        if st.session_state.view_mode == "editor":
            self._render_editor()
        else:
            self._render_preview()
    
    def _extract_text_from_docx(self, docx_file):
        """
        Extract text from a Word document.
        
        Args:
            docx_file: The uploaded Word document
            
        Returns:
            str: The extracted text content
        """
        try:
            # Load the document
            doc_bytes = docx_file.getvalue()
            
            doc = docx.Document(io.BytesIO(doc_bytes))
            full_text = []
            
            # Extract text from paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        table_count += 1
            
            result = "\n\n".join(full_text)
            return result
        except Exception as e:
            st.error(f"Error processing Word document: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            return f"Error processing document: {str(e)}"
    
    def _extract_text_from_pdf(self, pdf_file):
        """
        Extract text from a PDF file.
        
        Args:
            pdf_file: The uploaded PDF file
            
        Returns:
            str: The extracted text content
        """
        try:
            # Load the PDF
            pdf_bytes = pdf_file.getvalue()
            
            pdf_reader = PdfReader(io.BytesIO(pdf_bytes))
            full_text = []
            
            # Extract text from each page
            page_count = len(pdf_reader.pages)
            
            for page_num in range(page_count):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    full_text.append(text)
            
            result = "\n\n".join(full_text)
            return result
        except Exception as e:
            st.error(f"Error processing PDF file: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            return f"Error processing document: {str(e)}"
    
    def _render_editor(self):
        """
        Render the markdown editor.
        
        Displays a full-width text area for editing markdown content with
        syntax highlighting and proper styling.
        """
        # Markdown editor with improved styling
        new_content = st.text_area(
            "Edit your document here",
            value=st.session_state.markdown_content,
            height=500,  # Increased height for better editing experience
            label_visibility="collapsed",
            key="markdown_editor"
        )
        
        # Display word count
        word_count = len(new_content.split()) if new_content else 0
        char_count = len(new_content) if new_content else 0
        st.caption(f"Word count: {word_count} | Character count: {char_count}")
        
        # Update content if changed
        if new_content != st.session_state.markdown_content:
            # Save current state to history before updating
            self._save_to_history()
            st.session_state.markdown_content = new_content
            if self.on_content_change:
                self.on_content_change(new_content)
    
    def _render_preview(self):
        """
        Render the markdown preview.
        
        Displays a full-width preview of the rendered markdown content
        with proper styling and formatting.
        
        Implementation note:
        This method uses Streamlit's native expander component to contain the markdown content
        rather than custom HTML/CSS containers. This approach eliminates the excessive whitespace
        issues that can occur with custom containers and provides a cleaner, more consistent
        rendering of the markdown content.
        """
        # Add download button at the top
        if st.session_state.markdown_content:
            filename = st.session_state.current_file or "document.md"
            st.download_button(
                label="💾 Download",
                data=st.session_state.markdown_content,
                file_name=filename,
                mime="text/markdown",
            )
            
            # Use a single expander that's always expanded to contain the content
            # This helps eliminate whitespace issues by using Streamlit's native components
            with st.expander("", expanded=True):
                # Apply minimal styling to the content
                st.markdown(st.session_state.markdown_content)
        else:
            st.info("Your document preview will appear here. Switch to the Editor to start writing.")
    
    def _save_to_history(self):
        """Save current content to undo history."""
        # Only save if there's content and it's different from the last saved state
        if "markdown_content" in st.session_state and st.session_state.markdown_content:
            # Initialize history if not exists
            if "undo_history" not in st.session_state:
                st.session_state.undo_history = []
                
            # Add current content to history if different from last entry
            if not st.session_state.undo_history or st.session_state.markdown_content != st.session_state.undo_history[-1]:
                st.session_state.undo_history.append(st.session_state.markdown_content)
                
                # Limit history size
                if len(st.session_state.undo_history) > 10:  # Hardcoded limit to avoid potential issues
                    st.session_state.undo_history = st.session_state.undo_history[-10:]
    
    def _undo_last_change(self):
        """Undo the last change by restoring content from history."""
        if "undo_history" in st.session_state and st.session_state.undo_history:
            # Get the previous content (last item in history)
            previous_content = st.session_state.undo_history.pop()
            
            # Update the content
            st.session_state.markdown_content = previous_content
            
            # Notify about content change
            if self.on_content_change:
                self.on_content_change(previous_content)
            
            # Rerun to update the UI
            st.rerun()
    
    def get_content(self):
        """
        Get the current markdown content.
        
        Returns:
            str: The current markdown content
        """
        return st.session_state.markdown_content if "markdown_content" in st.session_state else ""
    
    def set_content(self, content, save_history=True):
        """
        Set the markdown content.
        
        Args:
            content (str): The new content
            save_history (bool): Whether to save the current content to history
        """
        if save_history and "markdown_content" in st.session_state:
            self._save_to_history()
            
        st.session_state.markdown_content = content
        
        if self.on_content_change:
            self.on_content_change(content) 